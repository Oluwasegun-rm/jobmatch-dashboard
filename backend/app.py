from __future__ import annotations

import os
import re
import logging
from fastapi import FastAPI, HTTPException, Query, UploadFile, File, Header
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response, PlainTextResponse
from pydantic import BaseModel
from typing import Any, Dict

from jobmatch.analyzer import analyze
from jobmatch.storage import save_analysis, fetch_recent, fetch_by_id, init_db, get_user_by_username, create_user, update_display_name, get_user_by_id, update_username, update_password_hash
from jobmatch.providers import remotive_client
from jobmatch.providers import usajobs_client
from jobmatch.providers.models import JobItem
from pypdf import PdfReader
try:
    import docx  # type: ignore
except Exception:  # pragma: no cover - optional dep resolution handled by requirements
    docx = None
from jobmatch.ai import resume_feedback as ai_resume_feedback
from jobmatch.ai import bullet_rewrites as ai_bullet_rewrites
from jobmatch.ai import cover_letter as ai_cover_letter
from jobmatch.config import load_config
from jobmatch.auth import hash_password, verify_password, create_token, decode_token, extract_bearer_token
from readability import Document  # type: ignore
from lxml import html as lxml_html  # type: ignore
import httpx
from urllib.parse import urlparse, parse_qs
import ipaddress


logging.basicConfig(level=logging.INFO, format='[%(asctime)s] %(levelname)s %(name)s: %(message)s')
log = logging.getLogger("jobmatch.app")

app = FastAPI(title="JobMatch AI Backend", version="0.1.0")

# CORS for frontend calls
def _parse_origins(val: str) -> list[str]:
    items: list[str] = []
    for raw in (val or "").split(","):
        t = raw.strip().strip('"').strip("'")
        if t:
            items.append(t)
    return items

origins_env = os.getenv("ALLOWED_ORIGINS", "*")
origin_regex_env = os.getenv("ALLOWED_ORIGIN_REGEX")
origins = _parse_origins(origins_env)

# If wildcard is configured, prefer a permissive regex to support allow_credentials
allow_origin_regex: str | None = None
if not origins or origins == ["*"]:
    allow_origin_regex = origin_regex_env or ".*"
    origins = []
else:
    allow_origin_regex = origin_regex_env or None

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_origin_regex=allow_origin_regex,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _origin_allowed(origin: str | None) -> bool:
    if not origin:
        return False
    if origin in origins:
        return True
    if allow_origin_regex:
        try:
            return re.match(allow_origin_regex, origin) is not None
        except Exception:
            return False
    return False


class PreflightCORS(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):  # type: ignore[override]
        origin = request.headers.get("origin")
        if request.method == "OPTIONS" and _origin_allowed(origin):
            headers = {
                "Access-Control-Allow-Origin": origin,
                "Access-Control-Allow-Credentials": "true",
                "Access-Control-Allow-Methods": "GET,POST,PUT,PATCH,DELETE,OPTIONS",
                "Access-Control-Allow-Headers": request.headers.get(
                    "access-control-request-headers", "Authorization, Content-Type, Accept, X-Requested-With"
                ),
                "Vary": "Origin",
            }
            return PlainTextResponse("", status_code=200, headers=headers)

        # Regular request: proceed and ensure headers are present when allowed
        response: Response = await call_next(request)
        if _origin_allowed(origin):
            response.headers["Access-Control-Allow-Origin"] = origin  # echo back
            response.headers["Access-Control-Allow-Credentials"] = "true"
            response.headers.setdefault("Vary", "Origin")
        return response


# Add our preflight handler last so it runs before others
app.add_middleware(PreflightCORS)


class AnalyzeRequest(BaseModel):
    resume_text: str
    job_text: str


@app.get("/health")
def health() -> Dict[str, str]:
    return {"status": "ok"}


@app.on_event("startup")
def _startup() -> None:
    # Ensure DB exists
    init_db()
    # Log sanitized environment/config for debugging
    try:
        from jobmatch.config import load_config as _load

        cfg = _load()
        def _mask(s: str | None) -> str:
            if not s:
                return "<empty>"
            if len(s) <= 8:
                return "****"
            return f"****{s[-4:]}"
        log.info(
            "Startup config: OPENAI_ENABLED=%s OPENAI_MODEL=%s OPENAI_API_KEY=%s DB_PATH=%s ALLOWED_ORIGINS=%s ORIGIN_REGEX=%s",
            bool(cfg.openai_enabled), cfg.openai_model, _mask(cfg.openai_api_key), cfg.db_path, ",".join(origins) or "* (regex)", allow_origin_regex or "<none>",
        )
    except Exception as e:
        log.warning("Failed to log startup config: %s", type(e).__name__)


@app.post("/analyze")
def analyze_endpoint(req: AnalyzeRequest, authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    # Basic input size guardrails
    if len(req.resume_text) > 50_000 or len(req.job_text) > 50_000:
        raise HTTPException(status_code=413, detail="Input too large (limit 50k chars)")
    result = analyze(req.resume_text, req.job_text)
    # Log meta of AI usage from the result (if present)
    try:
        data = result.to_dict()
        meta = data.get("meta", {})
        matched_ct = len(data.get("matched_skills", []) or [])
        missing_ct = len(data.get("missing_skills", []) or [])
        log.info(
            "Analyze: score=%s ai_used=%s narrative_source=%s matches=%s missing=%s resume_len=%s job_len=%s",
            data.get("score"), meta.get("ai_used"), meta.get("narrative_source"), matched_ct, missing_ct, len(req.resume_text or ""), len(req.job_text or ""),
        )
        return {"ok": True, "result": data}
    except Exception:
        return {"ok": True, "result": result.to_dict()}


class SaveRequest(BaseModel):
    resume_text: str
    job_text: str
    score: int
    matched_skills: list[str]
    missing_skills: list[str]
    job_source: str | None = None
    job_url: str | None = None
    job_title: str | None = None
    job_company: str | None = None


@app.post("/save")
def save_endpoint(req: SaveRequest, authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    user_id: int | None = None
    tok = extract_bearer_token(authorization)
    if tok:
        payload = decode_token(tok)
        if payload:
            try:
                user_id = int(payload.get("sub"))
            except Exception:
                user_id = None
    rid = save_analysis(
        resume_text=req.resume_text,
        job_text=req.job_text,
        score=req.score,
        matched_skills=req.matched_skills,
        missing_skills=req.missing_skills,
        job_source=req.job_source,
        job_url=req.job_url,
        job_title=req.job_title,
        job_company=req.job_company,
        user_id=user_id,
    )
    return {"ok": True, "id": rid}


@app.get("/recent")
def recent(limit: int = 10, authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    user_id: int | None = None
    tok = extract_bearer_token(authorization)
    if tok:
        payload = decode_token(tok)
        if payload:
            try:
                user_id = int(payload.get("sub"))
            except Exception:
                user_id = None
    rows = fetch_recent(limit=limit, user_id=user_id)
    return {"ok": True, "results": rows}


# Jobs endpoints (Providers)

@app.get("/jobs/categories")
async def jobs_categories(source: str = Query("remotive")) -> Dict[str, Any]:
    if source == "remotive":
        cats = await remotive_client.fetch_categories()
        return {"ok": True, "results": [{"id": c.id, "name": c.name} for c in cats]}
    elif source == "usajobs":
        # No category list for USAJOBS in this app
        return {"ok": True, "results": []}
    else:
        raise HTTPException(status_code=400, detail="Unsupported jobs source")


@app.get("/jobs/search")
async def jobs_search(
    query: str = "",
    location: str | None = None,
    category: str | None = None,
    source: str = Query("remotive"),
    page: int = 1,
    per_page: int = 50,
    job_type: str | None = None,
) -> Dict[str, Any]:
    if source not in {"remotive", "usajobs"}:
        raise HTTPException(status_code=400, detail="Unsupported jobs source")
    # Bounds for pagination
    if page <= 0:
        page = 1
    per_page = max(10, min(per_page, 100))
    jobs: list[JobItem] = []
    if source == "remotive":
        # Pull enough items from provider so local pagination can fulfill per_page request
        # Fetch a cushion (2x current page size) up to provider cap to improve UX
        fetch_limit = max(100, min(500, per_page * max(1, page) * 2))
        try:
            jobs = await remotive_client.fetch_jobs(query=query, category=category, limit=fetch_limit)  # type: ignore[arg-type]
        except TypeError:
            # Backward compatibility with test doubles that don't accept 'limit'
            jobs = await remotive_client.fetch_jobs(query=query, category=category)
    else:
        # USAJOBS: provider-side pagination
        if not usajobs_client.is_enabled():
            raise HTTPException(status_code=503, detail="USAJOBS provider not configured")
        jobs = await usajobs_client.fetch_jobs(query=query, location=location or "United States", page=page, per_page=per_page)
    loc = (location or "").strip().lower()
    if loc:
        jobs = [j for j in jobs if loc in (j.location or "").lower()]
    jt = (job_type or "").strip().lower()
    if jt:
        # Accept exact or contains for internship etc.
        def _jt_ok(x: str | None) -> bool:
            s = (x or "").lower()
            if not s:
                return False
            if s == jt:
                return True
            # Soft contains for common synonyms
            if jt in {"intern", "internship", "internships"} and ("intern" in s):
                return True
            return jt in s
        jobs = [j for j in jobs if _jt_ok(j.job_type)]

    # Additional token-based filtering to improve matches
    tokens = [t for t in query.lower().split() if t]
    if tokens:
        def _hay(j: JobItem) -> str:
            parts = [j.title or "", j.company or "", j.location or "", j.description or ""]
            if j.tags:
                parts.extend(j.tags)
            return " \n ".join(parts).lower()
        all_matches = [j for j in jobs if all(tok in _hay(j) for tok in tokens)]
        if all_matches:
            jobs = all_matches
        else:
            # If strict 'all tokens' yields none, fall back to 'any token'
            any_matches = [j for j in jobs if any(tok in _hay(j) for tok in tokens)]
            jobs = any_matches or jobs

    total = len(jobs)
    start = (page - 1) * per_page
    end = start + per_page
    results = jobs[start:end] if source == "remotive" else jobs
    return {
        "ok": True,
        "page": page,
        "per_page": per_page,
        "total": total,
        "results": [
            {
                "id": j.id,
                "source": j.source,
                "title": j.title,
                "company": j.company,
                "location": j.location,
                "url": j.url,
                "posted_at": j.posted_at,
                "description": j.description,
                "job_type": j.job_type,
                "tags": j.tags,
                "salary": j.salary,
            }
            for j in results
        ],
    }


@app.get("/jobs/providers")
def jobs_providers() -> Dict[str, Any]:
    providers = [{"id": "remotive", "name": "Remotive"}]
    if usajobs_client.is_enabled():
        providers.append({"id": "usajobs", "name": "United States (USAJOBS)"})
    return {"ok": True, "results": providers}


@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)) -> Dict[str, Any]:
    """Accept a resume file and extract text.

    Supports .pdf, .docx, .txt. Returns extracted text for client-side analysis.
    """
    name = (file.filename or "").lower()
    content = await file.read()
    # Basic size limit: 8 MB
    if len(content) > 8 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (limit 8MB)")

    extracted = ""
    try:
        if name.endswith(".pdf"):
            # Use pypdf
            import io

            reader = PdfReader(io.BytesIO(content))
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    continue
            extracted = "\n".join(parts)
        elif name.endswith(".docx"):
            if docx is None:
                raise HTTPException(status_code=500, detail="DOCX support not available")
            import io

            doc = docx.Document(io.BytesIO(content))  # type: ignore[attr-defined]
            extracted = "\n".join(p.text for p in doc.paragraphs)
        elif name.endswith(".txt"):
            # Try utf-8 then latin-1
            try:
                extracted = content.decode("utf-8")
            except Exception:
                extracted = content.decode("latin-1", errors="ignore")
        else:
            raise HTTPException(status_code=415, detail="Unsupported file type. Use PDF, DOCX, or TXT.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")

    if not extracted.strip():
        raise HTTPException(status_code=422, detail="No text extracted from resume")

    return {"ok": True, "text": extracted}


class FeedbackRequest(BaseModel):
    resume_text: str
    job_text: str | None = None


@app.post("/ai/feedback")
def ai_feedback(req: FeedbackRequest) -> Dict[str, Any]:
    # Guard input length; keep this much for AI context
    resume_txt = (req.resume_text or "").strip()
    job_txt = (req.job_text or "").strip() if req.job_text else None
    if len(resume_txt) > 40_000 or (job_txt and len(job_txt) > 40_000):
        raise HTTPException(status_code=413, detail="Input too large (limit 40k chars)")

    cfg = load_config()
    suggestions: list[str] = []
    missing_keywords: list[str] = []
    rewrites: list[dict] = []
    if cfg.openai_enabled and cfg.openai_api_key:
        # Parse missing keywords from AI resume feedback if included
        try:
            # Reuse ai_resume_feedback to get suggestions; we also attempt to fetch missing keywords via a direct call
            # The function returns suggestions only for backward compatibility; we make a parallel lightweight call
            suggestions = ai_resume_feedback(resume_txt, job_txt)
        except Exception:
            suggestions = []
        # Try a few bullet rewrites as well
        rewrites = ai_bullet_rewrites(resume_txt, job_txt, max_items=3)

    # Heuristic fallback if AI disabled or no output
    if not suggestions:
        basic: list[str] = []
        t = resume_txt
        if len(t) < 400:
            basic.append("Add more detail. Aim for a concise 1–2 page resume with highlights.")
        if "responsible for" in t.lower():
            basic.append("Replace weak phrasing like 'responsible for' with stronger verbs (Led, Spearheaded).")
        if not any(ch.isdigit() for ch in t):
            basic.append("Quantify impact by adding metrics (%, $, #).")
        if not any(b in t for b in ["•", "- ", "* "]):
            basic.append("Use concise bullet points for readability.")
        suggestions = basic[:6]
    # Baseline missing skills from parser vs job
    try:
        from jobmatch.config import load_config as _load
        from jobmatch.parser import extract_skills
        from jobmatch.scorer import evaluate
        cfg2 = _load()
        rs = extract_skills(resume_txt, cfg2.alias_map)
        js = extract_skills(job_txt or "", cfg2.alias_map)
        missing_keywords = sorted(list(evaluate(js, rs).missing_skills))[:10]
    except Exception:
        missing_keywords = []

    return {"ok": True, "suggestions": suggestions, "missing_keywords": missing_keywords, "rewrites": rewrites}


class CoverLetterRequest(BaseModel):
    resume_text: str
    job_text: str
    tone: str | None = "professional"


@app.post("/ai/cover-letter")
def ai_cover_letter_endpoint(req: CoverLetterRequest) -> Dict[str, Any]:
    if len(req.resume_text or "") > 40_000 or len(req.job_text or "") > 40_000:
        raise HTTPException(status_code=413, detail="Input too large (limit 40k chars)")
    cfg = load_config()
    text = ""
    if cfg.openai_enabled and cfg.openai_api_key:
        text = ai_cover_letter(req.resume_text, req.job_text, (req.tone or "professional"))
    if not text:
        # Fallback template
        text = (
            "Dear Hiring Team,\n\nI’m excited to apply for this role. My background closely aligns with the job description, "
            "and I have hands-on experience in the listed technologies. I consistently focus on measurable outcomes, "
            "collaboration, and clear communication. I’d welcome the chance to discuss how I can contribute.\n\nSincerely,\nYour Name"
        )
    return {"ok": True, "cover_letter": text}


@app.get("/analysis/{analysis_id}")
def get_analysis(analysis_id: int) -> Dict[str, Any]:
    row = fetch_by_id(analysis_id)
    if not row:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"ok": True, "result": row}


# --- Job extraction from URL ---

class ExtractJobRequest(BaseModel):
    url: str


def _is_private_host(host: str) -> bool:
    try:
        ip = ipaddress.ip_address(host)
        return ip.is_private or ip.is_loopback or ip.is_link_local
    except ValueError:
        # Not an IP literal; basic denylist
        low = host.lower()
        if low in {"localhost"}:
            return True
        if low.endswith(".local") or low.endswith(".internal"):
            return True
        return False


@app.post("/extract-job")
def extract_job(req: ExtractJobRequest) -> Dict[str, Any]:
    raw = (req.url or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="url is required")
    try:
        p = urlparse(raw)
    except Exception:
        raise HTTPException(status_code=400, detail="invalid url")
    if p.scheme not in {"http", "https"} or not p.netloc:
        raise HTTPException(status_code=400, detail="invalid url")
    if _is_private_host(p.hostname or ""):
        raise HTTPException(status_code=400, detail="disallowed host")

    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; JobMatchBot/1.0; +https://github.com/Oluwasegun-rm/jobmatch-dashboard)",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    # Special-case Greenhouse embed links: rewrite to canonical job page
    if (p.hostname or "").endswith("greenhouse.io") and "/embed/job_app" in (p.path or ""):
        try:
            qs = parse_qs(p.query)
            company = (qs.get("for") or [None])[0]
            token = (qs.get("token") or [None])[0]
            if company and token:
                raw = f"https://boards.greenhouse.io/{company}/jobs/{token}"
                p = urlparse(raw)
                # Set a referer to reduce 404s on some hosts
                headers["Referer"] = "https://boards.greenhouse.io/"
        except Exception:
            pass
    try:
        with httpx.Client(timeout=10.0, follow_redirects=True, headers=headers) as client:
            resp = client.get(raw)
            resp.raise_for_status()
            content = resp.text or ""
            if len(content) > 4_000_000:
                raise HTTPException(status_code=413, detail="page too large")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"failed to fetch url: {type(e).__name__}")

    # Readability extraction
    try:
        doc = Document(content)
        title = (doc.short_title() or "").strip()
        summary_html = doc.summary(html_partial=False) or ""
        root = lxml_html.fromstring(summary_html)
        text = root.text_content().strip()
    except Exception:
        title, text = "", ""

    # Fallbacks if too short
    if len(text) < 240:
        try:
            root_full = lxml_html.fromstring(content)
            # Common selectors
            candidates = root_full.cssselect("#jobDescription, .job-description, .description, [itemprop=description], article")
            for el in candidates:
                t = el.text_content().strip()
                if len(t) > len(text):
                    text = t
            if not title:
                title_nodes = root_full.cssselect("title")
                if title_nodes:
                    title = (title_nodes[0].text or "").strip()
            if not text:
                # meta description
                metas = root_full.cssselect('meta[name="description"]')
                if metas:
                    text = (metas[0].get("content") or "").strip()
        except Exception:
            pass

    if not (text and text.strip()):
        raise HTTPException(status_code=422, detail="could not extract job text; paste manually")

    return {"ok": True, "title": title, "text": text, "source_url": raw}


# --- Auth Endpoints ---

class SignupRequest(BaseModel):
    username: str
    password: str
    display_name: str | None = None


@app.post("/auth/signup")
def auth_signup(req: SignupRequest) -> Dict[str, Any]:
    if not req.username or not req.password:
        raise HTTPException(status_code=400, detail="username and password required")
    if get_user_by_username(req.username):
        raise HTTPException(status_code=409, detail="username is taken")
    ph = hash_password(req.password)
    uid = create_user(req.username, ph, req.display_name or None, is_admin=False)
    token = create_token(uid, req.username, req.display_name or req.username)
    return {"ok": True, "token": token, "user": {"id": uid, "username": req.username, "display_name": req.display_name or req.username}}


class LoginRequest(BaseModel):
    username: str
    password: str


@app.post("/auth/login")
def auth_login(req: LoginRequest) -> Dict[str, Any]:
    u = get_user_by_username(req.username)
    if not u or not verify_password(req.password, u["password_hash"]):
        raise HTTPException(status_code=401, detail="invalid credentials")
    token = create_token(u["id"], u["username"], u.get("display_name") or u["username"])
    return {"ok": True, "token": token, "user": {"id": u["id"], "username": u["username"], "display_name": u.get("display_name") or u["username"]}}


@app.get("/auth/me")
def auth_me(authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    tok = extract_bearer_token(authorization)
    if not tok:
        raise HTTPException(status_code=401, detail="missing token")
    p = decode_token(tok)
    if not p:
        raise HTTPException(status_code=401, detail="invalid token")
    try:
        uid = int(p.get("sub"))
    except Exception:
        raise HTTPException(status_code=401, detail="invalid token")
    user = get_user_by_id(uid)
    if not user:
        raise HTTPException(status_code=404, detail="user not found")
    return {"ok": True, "user": {"id": user["id"], "username": user["username"], "display_name": user.get("display_name") or user["username"]}}


class DisplayNameRequest(BaseModel):
    display_name: str


@app.post("/auth/display-name")
def auth_display_name(req: DisplayNameRequest, authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    tok = extract_bearer_token(authorization)
    p = decode_token(tok or "")
    if not p:
        raise HTTPException(status_code=401, detail="invalid token")
    uid = int(p.get("sub"))
    update_display_name(uid, req.display_name)
    return {"ok": True}


class ChangeUsernameRequest(BaseModel):
    username: str


@app.post("/auth/change-username")
def auth_change_username(req: ChangeUsernameRequest, authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    tok = extract_bearer_token(authorization)
    p = decode_token(tok or "")
    if not p:
        raise HTTPException(status_code=401, detail="invalid token")
    uid = int(p.get("sub"))
    if not req.username or len(req.username.strip()) < 3:
        raise HTTPException(status_code=400, detail="invalid username")
    # Check if taken
    existing = get_user_by_username(req.username.strip())
    if existing and int(existing["id"]) != uid:
        raise HTTPException(status_code=409, detail="username is taken")
    try:
        update_username(uid, req.username.strip())
    except Exception:
        raise HTTPException(status_code=409, detail="username is taken")
    # Issue a refreshed token with updated username
    me = get_user_by_id(uid)
    token = create_token(uid, req.username.strip(), (me.get("display_name") if me else None) or req.username.strip())
    return {"ok": True, "token": token, "user": {"id": uid, "username": req.username.strip(), "display_name": (me.get("display_name") if me else None) or req.username.strip()}}


class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str


@app.post("/auth/change-password")
def auth_change_password(req: ChangePasswordRequest, authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    tok = extract_bearer_token(authorization)
    p = decode_token(tok or "")
    if not p:
        raise HTTPException(status_code=401, detail="invalid token")
    uid = int(p.get("sub"))
    u = get_user_by_id(uid)
    if not u or not verify_password(req.current_password or "", u["password_hash"]):
        raise HTTPException(status_code=401, detail="invalid current password")
    if not req.new_password or len(req.new_password) < 6:
        raise HTTPException(status_code=400, detail="new password too short")
    ph = hash_password(req.new_password)
    update_password_hash(uid, ph)
    return {"ok": True}
